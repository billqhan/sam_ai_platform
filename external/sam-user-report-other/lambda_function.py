import os
import json
import boto3
import logging
from datetime import datetime, timezone
from urllib.parse import unquote_plus
from io import BytesIO
from docx import Document
from docx.shared import Pt
import re # Import the re module for regular expressions

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
s3_client = boto3.client('s3')
agent = boto3.client("bedrock-agent-runtime")

# Configuration - Updated for Bedrock Agents
DESTINATION_BUCKET = os.environ.get("DESTINATION_BUCKET", "m-sam-opportunity-responses")
AGENT_ID = os.environ["AGENT_ID"]  # Your Bedrock Agent ID
AGENT_ALIAS_ID = os.environ["AGENT_ALIAS_ID"]  # Your Bedrock Agent Alias ID
OUTPUT_FORMATS = os.environ.get('OUTPUT_FORMATS', 'txt,docx').split(',')

def lambda_handler(event, context):
    '''
    Lambda function triggered by S3 object creation.
    Processes JSON opportunity files and generates response templates using Bedrock Agent.
    Only processes files in YYYY-MM-DD/matches/ folder pattern.
    '''

    try:
        # Parse S3 event
        for record in event['Records']:
            # Get bucket and object key from the event
            source_bucket = record['s3']['bucket']['name']
            object_key = unquote_plus(record['s3']['object']['key'])

            logger.info(f"Processing file: {object_key} from bucket: {source_bucket}")

            # Skip if not a JSON file
            if not object_key.lower().endswith('.json'):
                logger.info(f"Skipping non-JSON file: {object_key}")
                continue

            # Check if file is in the correct folder pattern: YYYY-MM-DD/matches/
            if not is_valid_matches_path(object_key):
                logger.info(f"Skipping file not in matches folder: {object_key}")
                continue

            # Read JSON file from S3
            json_data = read_json_from_s3(source_bucket, object_key)

            # Generate response template using Bedrock Agent
            response_template = generate_response_template(json_data)

            # Save to destination bucket
            save_response_to_s3(response_template, json_data, object_key, formats=OUTPUT_FORMATS)

            logger.info(f"Successfully processed {object_key}")

    except Exception as e:
        logger.error(f"Error processing Lambda event: {str(e)}")
        raise e

    return {
        'statusCode': 200,
        'body': json.dumps('Processing completed successfully')
    }

def is_valid_matches_path(object_key):
    '''
    Check if the object key matches the pattern: YYYY-MM-DD/matches/*.json
    Returns True if the path is valid, False otherwise.
    '''
    import re
    
    # Pattern to match YYYY-MM-DD/matches/ folder structure
    # This regex matches: 4 digits, dash, 2 digits, dash, 2 digits, /matches/
    date_matches_pattern = r'^\d{4}-\d{2}-\d{2}/matches/'
    
    if re.match(date_matches_pattern, object_key):
        logger.info(f"Valid matches path detected: {object_key}")
        return True
    else:
        logger.info(f"Path does not match YYYY-MM-DD/matches/ pattern: {object_key}")
        return False

def read_json_from_s3(bucket, key):
    '''Read and parse JSON file from S3'''
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read().decode('utf-8')
        return json.loads(content)
    except Exception as e:
        logger.error(f"Error reading JSON from S3: {str(e)}")
        raise e

def call_agent(agent_id, agent_alias_id, prompt):
    '''Call Bedrock Agent with the given prompt'''
    resp = agent.invoke_agent(
        agentId=agent_id,
        agentAliasId=agent_alias_id,
        sessionId="sam-batch",
        inputText=prompt,
        enableTrace=False,
        endSession=False
    )
    full_text = ""
    for event in resp.get("completion", []):
        if "chunk" in event:
            full_text += event["chunk"]["bytes"].decode("utf-8")
    return full_text

def generate_response_template(json_data):
    '''Use Amazon Bedrock Agent to generate the response template'''

    # Construct the prompt for Bedrock Agent
    prompt_text = f'''You are a government contracting response assistant. Based on the following JSON opportunity record, create two distinct outputs:

JSON INPUT:
{json.dumps(json_data, indent=2)}

REQUIREMENTS:

Create exactly two sections:

**Section 1: Human-Readable Summary (For Sender)**
- Start with a one-line summary that includes the opportunity essence and match score
- Include opportunity overview with key details (solicitation number, title, agency, dates)
- Provide a plain-English description of what the government is seeking
- List place of performance and government point of contact
- Include match assessment details (score and rationale)
- List company data used in the match

**Section 2: Draft Response Template (Formal Email to Government)**
- Create a formal business letter format
- Use proper addressing with multiple lines for readability
- Include subject line with solicitation number
- Write in professional, sentence-format paragraphs (no bullet points)
- Include sections for: Company Information, Statement of Capability, Past Performance and References, Conclusion
- Preserve placeholders like [Your Company Name], [UEI], [CAGE Code] for missing company data
- In the Conclusion, incorporate and rewrite the rationale from the JSON into formal, confident language
- Do NOT use the word "section" in any headings within Section 2
- Keep all content professional and government-appropriate

IMPORTANT: 
- Do not include any hidden reasoning markers (like <reasoning>, , "Let's craft", or explanations).
- Only use information present in the JSON - do not invent details
- Maintain formal tone throughout Section 2
- Keep Section 1 conversational and easy to scan
- Ensure the match score is prominently displayed in Section 1 summary line
- Do not include any internal reasoning, thinking, or meta-commentary in your response
- Output must begin directly with "## Section 1:" and continue with final user‑readable text only.
- Never include words like "Based on", "Let's", "Thinking", or meta explanations.'''

    try:
        logger.info(f"Using Bedrock Agent: {AGENT_ID}")
        
        # Call Bedrock Agent
        generated_text = call_agent(AGENT_ID, AGENT_ALIAS_ID, prompt_text)

        # Clean up AI thinking traces and meta-commentary
        cleaned_text = clean_ai_thinking_traces(generated_text)

        logger.info(f"Successfully generated response template using Bedrock Agent {AGENT_ID}")
        return cleaned_text

    except Exception as e:
        logger.error(f"Error calling Bedrock Agent {AGENT_ID}: {str(e)}")
        raise e


def clean_ai_thinking_traces(text):
    '''Remove AI internal thinking, reasoning, and meta-commentary from the output'''
    
    import re
    
    # Remove common AI thinking patterns
    thinking_patterns = [
        r"Let me think about this.*?\n",
        r"I need to.*?\n",
        r"First, I'll.*?\n",
        r"Now I'll.*?\n",
        r"Looking at this.*?\n",
        r"Based on my analysis.*?\n",
        r"I should.*?\n",
        r"I will.*?\n",
        r"I can see.*?\n",
        r"It appears.*?\n",
        r"I notice.*?\n",
        r"Let me analyze.*?\n",
        r"I'm going to.*?\n",
        r"Here's what I.*?\n",
        r"I understand.*?\n",
        r"From what I can see.*?\n",
        r"Looking at the JSON.*?\n",
        r"Based on the provided information.*?\n",
        r"Let's craft.*?\n",
        r"Here's how.*?\n",
        r"I'll create.*?\n"
    ]
    
    # Remove thinking patterns (case insensitive)
    cleaned_text = text
    for pattern in thinking_patterns:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove text before the first "## Section 1" if it exists
    section_match = re.search(r'## Section 1:', cleaned_text, re.IGNORECASE)
    if section_match:
        cleaned_text = cleaned_text[section_match.start():]
    
    # Remove meta-instructions that look like bullet point lists
    # This catches patterns like "Section 1 details:" followed by bullet points
    meta_instruction_patterns = [
        r'Section \d+ details?:.*?(?=##|\Z)',  # "Section 1 details:" until next section or end
        r'Section \d+ requirements?:.*?(?=##|\Z)',
        r'Here are the details?:.*?(?=##|\Z)',
        r'Details for Section \d+:.*?(?=##|\Z)',
        r'Requirements for Section \d+:.*?(?=##|\Z)',
        r'- [A-Z][^.]*\. [A-Z][^.]*\.[^\n]*\n',  # Bullet points that look like instructions
        r'- [A-Z][^:]*:[^\n]*\n',  # Bullet points with colons (like "- One-line summary includes:")
    ]
    
    for pattern in meta_instruction_patterns:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)

    # Remove any remaining meta-commentary in parentheses or brackets
    cleaned_text = re.sub(r'\([^)]*thinking[^)]*\)', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'\[[^\]]*reasoning[^\]]*\]', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'<reasoning>.*?</reasoning>', '', cleaned_text, flags=re.IGNORECASE | re.DOTALL)

    # Remove standalone instruction lines that start with dashes
    instruction_lines = [
        r'^- [A-Z][^.]*includes.*?\n',  # Lines like "- One-line summary includes..."
        r'^- [A-Z][^.]*details.*?\n',   # Lines like "- Include opportunity overview..."
        r'^- [A-Z][^.]*provide.*?\n',   # Lines like "- Provide a plain-English..."
        r'^- [A-Z][^.]*list.*?\n',    # Lines like "- List place of performance..."
        r'^- [A-Z][^.]*create.*?\n',    # Lines like "- Create a formal business..."
        r'^- [A-Z][^.]*use.*?\n',    # Lines like "- Use proper addressing..."
        r'^- [A-Z][^.]*write.*?\n',    # Lines like "- Write in professional..."
        r'^- [A-Z][^.]*include.*?\n',   # Lines like "- Include sections for..."
    ]
    
    for pattern in instruction_lines:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE)


    # Remove excessive whitespace and clean up formatting
    cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)  # Max 2 consecutive newlines
    cleaned_text = cleaned_text.strip()
    
    # Log if significant cleaning occurred
    original_length = len(text)
    cleaned_length = len(cleaned_text)
    if original_length - cleaned_length > 100:
        logger.info(f"Cleaned AI thinking traces: removed {original_length - cleaned_length} characters")
    
    return cleaned_text



def add_markdown_paragraph(doc, line):
    """
    Convert Markdown bold (**text**) into actual bold text
    and add line to Word doc as a paragraph.
    """
    para = doc.add_paragraph()
    bold_pattern = r"\*\*(.*?)\*\*"
    
    last_end = 0
    for match in re.finditer(bold_pattern, line):
        # Normal text before bold
        if match.start() > last_end:
            para.add_run(line[last_end:match.start()])
        
        # Bold text
        bold_text = match.group(1)
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        
        last_end = match.end()
    
    # Remaining text after last bold
    if last_end < len(line):
        para.add_run(line[last_end:])

def generate_docx(response_content, metadata):
    """
    Generate a DOCX file with proper formatting and page breaks.
    """
    doc = Document()
    ts = metadata["timestamp"]

    # Title Page
    doc.add_heading('Response Template', level=1)
    doc.add_paragraph(f"Generated: {ts}")
    doc.add_paragraph(f"Source File: {metadata['source_file']}")
    doc.add_paragraph(f"Solicitation: {metadata['solicitation']}")
    doc.add_paragraph(f"Match Score: {metadata['match_score']}")
    doc.add_paragraph(f"Agent: {metadata['agent']}")
    doc.add_paragraph("---")

    # Response content, line by line
    for line in response_content.splitlines():
        line = line.strip()
        if not line:
            continue

        # Insert page breaks before Section 1 and Section 2
        if line.lower().startswith("## section 1"):
            doc.add_page_break()
            doc.add_heading("Section 1: Human-Readable Summary (For Sender)", level=2)
            continue
        elif line.lower().startswith("## section 2"):
            doc.add_page_break()
            doc.add_heading("Section 2: Draft Response Template (Formal Email to Government)", level=2)
            continue

        # Convert markdown bold (**text**) into real bold runs
        add_markdown_paragraph(doc, line)

    return doc

def save_response_to_s3(response_content, json_data, original_key, formats=None):
    """
    Save the generated response template to S3 in multiple formats.
    Formats determined by OUTPUT_FORMATS environment variable or passed parameter.
    """
    if formats is None:
        formats = OUTPUT_FORMATS  # Use environment variable

    try:
        solicitation_number = json_data.get('solicitationNumber', 'UNKNOWN')

        # Extract prefix (parent folder path) from the original key
        prefix = os.path.dirname(original_key)
        if prefix in ["", "."]:
            prefix = ""  # no subfolder
        else:
            prefix = prefix + "/"  # ensure trailing slash

        # Common filename stem
        ## base_filename = original_key.replace('.json', '').split('/')[-1]
        base_filename = os.path.splitext(os.path.basename(original_key))[0]
        output_basename = f"{solicitation_number}_response_template"
        
        # Metadata
        timestamp = datetime.utcnow().isoformat()
        header = f"""# Response Template Generated: {timestamp}
# Source File: {original_key}
# Solicitation: {solicitation_number}
# Match Score: {json_data.get('score', 'N/A')}
# Agent Used: {AGENT_ID}

---
"""

        # --- Save as TXT (Default) ---
        if "txt" in formats:
            txt_filename = f"{prefix}{output_basename}.txt"
            full_content = header + response_content
            s3_client.put_object(
                Bucket=DESTINATION_BUCKET,
                Key=txt_filename,
                Body=full_content.encode('utf-8'),
                ContentType='text/plain',
                Metadata={
                    'source-file': original_key,
                    'solicitation-number': solicitation_number,
                    'generated-timestamp': timestamp,
                    'match-score': str(json_data.get('score', 'N/A')),
                    'agent-used': AGENT_ID
                }
            )
            logger.info(f"Saved TXT response template: {txt_filename}")
        
        # --- Save as DOCX ---
        if "docx" in formats:
            metadata = {
                "timestamp": timestamp,
                "source_file": original_key,
                "solicitation": solicitation_number,
                "match_score": str(json_data.get('score', 'N/A')),
                "agent": AGENT_ID,
            }
            doc = generate_docx(response_content, metadata)


            # Save DOCX in memory
            byte_stream = BytesIO()
            doc.save(byte_stream)
            byte_stream.seek(0)

            docx_filename = f"{prefix}{output_basename}.docx"
            s3_client.put_object(
                Bucket=DESTINATION_BUCKET,
                Key=docx_filename,
                Body=byte_stream.getvalue(),
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                Metadata={
                    'source-file': original_key,
                    'solicitation-number': solicitation_number,
                    'generated-timestamp': timestamp,
                    'match-score': str(json_data.get('score', 'N/A')),
                    'agent-used': AGENT_ID
                }
            )
            logger.info(f"Saved DOCX response template with formatting: {docx_filename}")

    except Exception as e:
        logger.error(f"Error saving to S3: {str(e)}")
        raise e