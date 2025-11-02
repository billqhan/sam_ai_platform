"""
Lambda function entry point for sam-produce-user-report.
Processes JSON opportunity files and generates response templates using Bedrock Agent.
Only processes files in YYYY-MM-DD/matches/ folder pattern.
"""

import os
import json
import boto3
import logging
from datetime import datetime, timezone
from urllib.parse import unquote_plus
from io import BytesIO
import re

# Try to import docx, but handle gracefully if not available
try:
    from docx import Document
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available, will generate RTF format instead")

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Configuration - Updated for Bedrock Knowledge Base
DESTINATION_BUCKET = os.environ.get("DESTINATION_BUCKET", "m-sam-opportunity-responses")
KNOWLEDGE_BASE_ID = os.environ.get("KNOWLEDGE_BASE_ID", "")
MODEL_ID = os.environ.get("MODEL_ID", "amazon.nova-pro-v1:0")
MAX_TOKENS = int(os.environ.get("MAX_TOKENS", "8000"))
TEMPERATURE = float(os.environ.get("TEMPERATURE", "0.1"))
BEDROCK_REGION = os.environ.get("BEDROCK_REGION", "us-east-1")

# Initialize AWS clients
s3_client = boto3.client('s3')
bedrock_client = boto3.client('bedrock-runtime', region_name=BEDROCK_REGION)
bedrock_agent_runtime = boto3.client('bedrock-agent-runtime', region_name=BEDROCK_REGION)
OUTPUT_FORMATS = os.environ.get('OUTPUT_FORMATS', 'txt,docx').split(',')

def lambda_handler(event, context):
    '''
    Lambda function triggered by S3 object creation.
    Processes JSON opportunity files and generates response templates using Bedrock Agent.
    Only processes files in YYYY-MM-DD/matches/ folder pattern.
    '''
    try:
        # Parse S3 event
        for record in event['Records']:
            # Get bucket and object key from the event
            source_bucket = record['s3']['bucket']['name']
            object_key = unquote_plus(record['s3']['object']['key'])
            
            logger.info(f"Processing file: {object_key} from bucket: {source_bucket}")
            
            # Skip if not a JSON file
            if not object_key.lower().endswith('.json'):
                logger.info(f"Skipping non-JSON file: {object_key}")
                continue
            
            # Check if file is in the correct folder pattern: YYYY-MM-DD/matches/
            if not is_valid_matches_path(object_key):
                logger.info(f"Skipping file not in matches folder: {object_key}")
                continue
            
            # Read JSON file from S3
            json_data = read_json_from_s3(source_bucket, object_key)
            
            # Generate response template using Bedrock Agent
            response_template = generate_response_template(json_data)
            
            # Save to destination bucket
            save_response_to_s3(response_template, json_data, object_key, formats=OUTPUT_FORMATS)
            
            logger.info(f"Successfully processed {object_key}")
            
    except Exception as e:
        logger.error(f"Error processing Lambda event: {str(e)}")
        raise e
    
    return {
        'statusCode': 200,
        'body': json.dumps('Processing completed successfully')
    }

def is_valid_matches_path(object_key):
    '''
    Check if the object key matches the pattern: YYYY-MM-DD/matches/*.json
    Returns True if the path is valid, False otherwise.
    '''
    # Pattern to match YYYY-MM-DD/matches/ folder structure
    # This regex matches: 4 digits, dash, 2 digits, dash, 2 digits, /matches/
    date_matches_pattern = r'^\d{4}-\d{2}-\d{2}/matches/'
    
    if re.match(date_matches_pattern, object_key):
        logger.info(f"Valid matches path detected: {object_key}")
        return True
    else:
        logger.info(f"Path does not match YYYY-MM-DD/matches/ pattern: {object_key}")
        return False

def read_json_from_s3(bucket, key):
    '''Read and parse JSON file from S3'''
    try:
        response = s3_client.get_object(Bucket=bucket, Key=key)
        content = response['Body'].read().decode('utf-8')
        return json.loads(content)
    except Exception as e:
        logger.error(f"Error reading JSON from S3: {str(e)}")
        raise e

def query_knowledge_base(query_text, max_results=10):
    '''Query the Bedrock Knowledge Base for company capability information'''
    if not KNOWLEDGE_BASE_ID:
        logger.warning("No KNOWLEDGE_BASE_ID configured, returning empty results")
        return []
    
    try:
        logger.info(f"Querying Knowledge Base {KNOWLEDGE_BASE_ID} with query: {query_text[:100]}...")
        
        # Query the knowledge base using bedrock-agent-runtime
        response = bedrock_agent_runtime.retrieve(
            knowledgeBaseId=KNOWLEDGE_BASE_ID,
            retrievalQuery={
                'text': query_text
            },
            retrievalConfiguration={
                'vectorSearchConfiguration': {
                    'numberOfResults': max_results
                }
            }
        )
        
        # Format results
        kb_results = []
        retrieval_results = response.get('retrievalResults', [])
        
        for i, result in enumerate(retrieval_results):
            content = result.get('content', {})
            location = result.get('location', {})
            metadata = result.get('metadata', {})
            
            # Extract text content
            text_content = content.get('text', '')
            
            # Create formatted result
            formatted_result = {
                'index': i,
                'title': metadata.get('title', f"Document {i+1}"),
                'snippet': text_content[:500] + "..." if len(text_content) > 500 else text_content,
                'source': location.get('s3Location', {}).get('uri', 'Unknown Source'),
                'metadata': metadata,
                'score': result.get('score', 0.0),
                'full_content': text_content
            }
            
            kb_results.append(formatted_result)
        
        logger.info(f"Successfully retrieved {len(kb_results)} results from Knowledge Base")
        return kb_results
        
    except Exception as e:
        logger.error(f"Error querying Knowledge Base: {str(e)}")
        return []

def call_bedrock_model(prompt):
    '''Call Bedrock model with the given prompt'''
    try:
        # Prepare the request body based on model type
        if 'claude' in MODEL_ID.lower():
            request_body = {
                "anthropic_version": "bedrock-2023-05-31",
                "max_tokens": MAX_TOKENS,
                "temperature": TEMPERATURE,
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ]
            }
        elif 'nova' in MODEL_ID.lower():
            # Nova-specific format
            request_body = {
                "messages": [
                    {
                        "role": "user",
                        "content": [{"text": prompt}]
                    }
                ],
                "inferenceConfig": {
                    "max_new_tokens": MAX_TOKENS,
                    "temperature": TEMPERATURE
                }
            }
        else:
            # Generic format for other models
            request_body = {
                "inputText": prompt,
                "textGenerationConfig": {
                    "maxTokenCount": MAX_TOKENS,
                    "temperature": TEMPERATURE
                }
            }
        
        # Call Bedrock model
        response = bedrock_client.invoke_model(
            modelId=MODEL_ID,
            body=json.dumps(request_body),
            contentType='application/json',
            accept='application/json'
        )
        
        # Parse response
        response_body = json.loads(response['body'].read())
        
        # Extract text based on model type
        if 'claude' in MODEL_ID.lower():
            if 'content' in response_body and len(response_body['content']) > 0:
                return response_body['content'][0]['text']
            else:
                logger.warning("Unexpected Claude response format")
                return ""
        elif 'nova' in MODEL_ID.lower():
            # Nova response format
            if 'output' in response_body and 'message' in response_body['output']:
                return response_body['output']['message']['content'][0]['text']
            else:
                logger.warning("Unexpected Nova response format")
                return str(response_body)
        else:
            # Generic format
            return response_body.get('results', [{}])[0].get('outputText', '')
            
    except Exception as e:
        logger.error(f"Error calling Bedrock model: {str(e)}")
        raise e

def get_nested_value(data, path, default='Unknown'):
    """Safely extract nested values from JSON using dot notation"""
    try:
        keys = path.split('.')
        value = data
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return default
        return value if value is not None else default
    except:
        return default

def clean_enhanced_description(description):
    """Clean enhanced description by removing unwanted headers and formatting"""
    if not description:
        return description
    
    # Remove BUSINESS SUMMARY headers and related formatting
    cleaned = description.replace('**BUSINESS SUMMARY:**', '')
    cleaned = cleaned.replace('BUSINESS SUMMARY:', '')
    
    # Remove other unwanted headers that might appear
    cleaned = cleaned.replace('**Purpose of the Solicitation:**', '')
    cleaned = cleaned.replace('Purpose of the Solicitation:', '')
    
    # Clean up extra whitespace and newlines
    cleaned = ' '.join(cleaned.split())
    
    return cleaned.strip()

def generate_section1_programmatically(json_data):
    """Generate Section 1 programmatically using actual JSON data"""
    
    # Extract data with proper handling of nested fields
    solicitation_number = json_data.get('solicitationNumber', 'Unknown')
    title = json_data.get('title', 'Unknown Title')
    agency = json_data.get('fullParentPathName', 'Unknown Agency')
    posted_date = json_data.get('postedDate', 'Unknown')
    response_deadline = json_data.get('responseDeadLine', 'Unknown')
    notice_type = json_data.get('type', 'Unknown')
    notice_id = json_data.get('noticeId', 'Unknown')
    ui_link = json_data.get('uiLink', 'Not available')
    
    # Handle nested point of contact fields
    poc_name = get_nested_value(json_data, 'pointOfContact.fullName', 'Unknown')
    poc_email = get_nested_value(json_data, 'pointOfContact.email', 'Unknown')
    poc_phone = get_nested_value(json_data, 'pointOfContact.phone', 'Not provided')
    
    # Handle nested place of performance fields
    city = get_nested_value(json_data, 'placeOfPerformance.city.name', 'Unknown')
    state = get_nested_value(json_data, 'placeOfPerformance.state.name', 'Unknown')
    country = get_nested_value(json_data, 'placeOfPerformance.country.name', 'Unknown')
    
    # Match assessment data
    score = json_data.get('score', 0)
    score_percentage = score * 100 if score else 0
    match_status = "Matched" if score >= 0.5 else "Not Matched"
    rationale = json_data.get('rationale', 'No rationale provided')
    
    # Company data arrays
    company_skills = json_data.get('company_skills', [])
    required_skills = json_data.get('opportunity_required_skills', [])
    past_performance = json_data.get('past_performance', [])
    kb_results_count = len(json_data.get('kb_retrieval_results', []))
    
    # Processing metadata
    input_key = json_data.get('input_key', 'Unknown')
    timestamp = json_data.get('timestamp', 'Unknown')
    enhanced_desc = json_data.get('enhanced_description', '')
    enhanced_desc_length = len(enhanced_desc)
    citations_count = len(json_data.get('citations', []))
    
    # Build Section 1 programmatically with proper spacing
    section1 = f"""**Section 1: Human-Readable Summary (For Sender)**

**Opportunity Overview**
- Solicitation Number: {solicitation_number}
- Title: {title}
- Agency: {agency}
- Posted Date: {posted_date}
- Response Deadline: {response_deadline}
- Notice Type: {notice_type}
- Notice ID: {notice_id}
- SAM.gov Link: {ui_link}

**What the Government Is Seeking**

{clean_enhanced_description(enhanced_desc)[:800] + '...' if len(enhanced_desc) > 800 else clean_enhanced_description(enhanced_desc)}

**Place of Performance & Point of Contact**
- Location: {city}, {state}, {country}
- Point of Contact: {poc_name}
- Email: {poc_email}
- Phone: {poc_phone}

**Match Assessment**
- Match Score: {score} ({score_percentage:.1f}%)
- Match Status: {match_status}
- Match Rationale: {rationale}

**Company Data Used in the Match**
- Company Skills: {', '.join(company_skills) if company_skills else 'None specified'}
- Opportunity Required Skills: {', '.join(required_skills) if required_skills else 'None specified'}
- Past Performance: {', '.join(past_performance) if past_performance else 'None specified'}
- Knowledge Base Results: {kb_results_count} relevant documents found

**Processing Information**
- Input File: {input_key}
- Processing Timestamp: {timestamp}
- Enhanced Description Length: {enhanced_desc_length} characters
- Citations Count: {citations_count}"""
    
    return section1

def generate_section2_programmatically(json_data):
    """Generate Section 2 programmatically using actual JSON data"""
    
    # Extract data for email template
    solicitation_number = json_data.get('solicitationNumber', 'UNKNOWN')
    title = json_data.get('title', 'Unknown Title')
    agency = json_data.get('fullParentPathName', 'Government Agency')
    poc_name = get_nested_value(json_data, 'pointOfContact.fullName', 'Contracting Officer')
    poc_email = get_nested_value(json_data, 'pointOfContact.email', 'contracting.officer@agency.gov')
    
    # Company skills and capabilities for content
    company_skills = json_data.get('company_skills', [])
    past_performance = json_data.get('past_performance', [])
    enhanced_desc = json_data.get('enhanced_description', '')
    rationale = json_data.get('rationale', 'We believe our qualifications and experience make us well-suited for this opportunity.')
    
    # Generate capability statement from company skills with proper grammar
    if company_skills:
        if len(company_skills) >= 3:
            capability_text = f"Our company possesses expertise in {company_skills[0]}, {company_skills[1]}, and {company_skills[2]}."
            if len(company_skills) > 3:
                additional_skills = company_skills[3:6]
                capability_text += f" Additionally, we have demonstrated experience with {', '.join(additional_skills)}."
        else:
            capability_text = f"Our company possesses expertise in {', '.join(company_skills)}."
    else:
        capability_text = "Our company possesses the technical expertise and experience required for this opportunity."
    
    # Generate past performance statement with proper grammar and capitalization
    if past_performance:
        if len(past_performance) == 1:
            # Ensure first letter is capitalized
            perf_item = past_performance[0].strip()
            if perf_item and not perf_item[0].isupper():
                perf_item = perf_item[0].upper() + perf_item[1:]
            performance_text = f"Our track record includes the following achievement: {perf_item}."
        elif len(past_performance) >= 2:
            # Ensure both items are properly capitalized
            perf_item1 = past_performance[0].strip()
            perf_item2 = past_performance[1].strip()
            if perf_item1 and not perf_item1[0].isupper():
                perf_item1 = perf_item1[0].upper() + perf_item1[1:]
            if perf_item2 and not perf_item2[0].isupper():
                perf_item2 = perf_item2[0].upper() + perf_item2[1:]
            performance_text = f"Our track record includes multiple achievements, including: {perf_item1}. We have also successfully completed: {perf_item2}."
        else:
            performance_text = f"Our past performance demonstrates success in {', '.join(past_performance[:3])}."
    else:
        performance_text = "We are establishing our performance record and are committed to delivering quality results on time and within budget."
    
    # Create a cleaner technical approach summary with proper capitalization
    if enhanced_desc:
        # Clean the description first, then extract first sentence
        cleaned_desc = clean_enhanced_description(enhanced_desc)
        first_sentence = cleaned_desc.split('.')[0] if '.' in cleaned_desc else cleaned_desc[:150]
        requirements_summary = first_sentence.strip()
        # Don't force lowercase - keep original capitalization for proper nouns like DARPA
    else:
        requirements_summary = "the requirements outlined in this solicitation"
    
    section2 = f"""**Section 2: Draft Response Template (Formal Email to Government)**

**TO:**
{poc_name}
{agency}
Email: {poc_email}

**FROM:**
[Your Name]
[Your Title]
[Your Company Name]
[Your Address]
[Your Phone]
[Your Email]

**DATE:** [Current Date]

**SUBJECT:** Response to Solicitation {solicitation_number} - {title}

Dear {poc_name},

[Your Company Name] respectfully submits this response to Solicitation {solicitation_number}, "{title}."

**Company Information**

[Your Company Name] is a qualified contractor with UEI [Your UEI] and CAGE Code [Your CAGE Code]. Our headquarters are located at [Your Company Address]. We maintain all requisite certifications and security clearances necessary for government contracting.

**Statement of Capability**

{capability_text} We have the technical expertise and resources necessary to successfully execute the requirements outlined in this solicitation.

**Past Performance and References**

{performance_text} Our commitment to excellence and customer satisfaction has been consistently demonstrated across all our engagements.

**Technical Approach**

We understand that this opportunity requires {requirements_summary}. Our approach will leverage our proven methodologies and experienced team to deliver a comprehensive solution that meets your objectives and exceeds expectations.

**Conclusion**

{rationale} We are committed to providing exceptional service and look forward to the opportunity to support your mission and contribute to the success of this important initiative.

Thank you for your consideration. Please contact me at [Your Phone Number] or [Your Email] for any additional information or clarification.

Respectfully submitted,

[Your Name]
[Your Title]
[Your Company Name]
[Your Phone Number]
[Your Email]

**Note:** Please replace all bracketed placeholders ([Your Name], [Your Company Name], etc.) with your actual company information before submitting."""
    
    return section2

def generate_response_template(json_data):
    """Generate response template using fully deterministic approach - no AI required"""
    
    try:
        logger.info("Using fully deterministic approach - generating both sections programmatically")
        
        # Generate both sections programmatically
        section1 = generate_section1_programmatically(json_data)
        section2 = generate_section2_programmatically(json_data)
        
        # Combine sections with separator
        full_response = f"{section1}\n\n---\n\n{section2}"
        
        logger.info("Successfully generated response template using deterministic approach")
        return full_response
        
    except Exception as e:
        logger.error(f"Error generating deterministic response template: {str(e)}")
        raise e

def clean_ai_thinking_traces(text):
    '''Remove AI internal thinking, reasoning, and meta-commentary from the output'''
    # Remove common AI thinking patterns
    thinking_patterns = [
        r"Let me think about this.*?\n",
        r"I need to.*?\n",
        r"First, I'll.*?\n",
        r"Now I'll.*?\n",
        r"Looking at this.*?\n",
        r"Based on my analysis.*?\n",
        r"I should.*?\n",
        r"I will.*?\n",
        r"I can see.*?\n",
        r"It appears.*?\n",
        r"I notice.*?\n",
        r"Let me analyze.*?\n",
        r"I'm going to.*?\n",
        r"Here's what I.*?\n",
        r"I understand.*?\n",
        r"From what I can see.*?\n",
        r"Looking at the JSON.*?\n",
        r"Based on the provided information.*?\n",
        r"Let's craft.*?\n",
        r"Here's how.*?\n",
        r"I'll create.*?\n"
    ]
    
    # Remove thinking patterns (case insensitive)
    cleaned_text = text
    for pattern in thinking_patterns:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove text before the first "## Section 1" if it exists
    section_match = re.search(r'## Section 1:', cleaned_text, re.IGNORECASE)
    if section_match:
        cleaned_text = cleaned_text[section_match.start():]
    
    # Remove meta-instructions that look like bullet point lists
    # This catches patterns like "Section 1 details:" followed by bullet points
    meta_instruction_patterns = [
        r'Section \d+ details?:.*?(?=##|\Z)',  # "Section 1 details:" until next section or end
        r'Section \d+ requirements?:.*?(?=##|\Z)',
        r'Here are the details?:.*?(?=##|\Z)',
        r'Details for Section \d+:.*?(?=##|\Z)',
        r'Requirements for Section \d+:.*?(?=##|\Z)',
        r'- [A-Z][^.]*\. [A-Z][^.]*\.[^\n]*\n',  # Bullet points that look like instructions
        r'- [A-Z][^:]*:[^\n]*\n',  # Bullet points with colons (like "- One-line summary includes:")
    ]
    
    for pattern in meta_instruction_patterns:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE | re.DOTALL)
    
    # Remove any remaining meta-commentary in parentheses or brackets
    cleaned_text = re.sub(r'\([^)]*thinking[^)]*\)', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'\[[^\]]*reasoning[^\]]*\]', '', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'<reasoning>.*?</reasoning>', '', cleaned_text, flags=re.IGNORECASE | re.DOTALL)
    
    # Remove standalone instruction lines that start with dashes
    instruction_lines = [
        r'^- [A-Z][^.]*includes.*?\n',  # Lines like "- One-line summary includes..."
        r'^- [A-Z][^.]*details.*?\n',   # Lines like "- Include opportunity overview..."
        r'^- [A-Z][^.]*provide.*?\n',   # Lines like "- Provide a plain-English..."
        r'^- [A-Z][^.]*list.*?\n',    # Lines like "- List place of performance..."
        r'^- [A-Z][^.]*create.*?\n',    # Lines like "- Create a formal business..."
        r'^- [A-Z][^.]*use.*?\n',    # Lines like "- Use proper addressing..."
        r'^- [A-Z][^.]*write.*?\n',    # Lines like "- Write in professional..."
        r'^- [A-Z][^.]*include.*?\n',   # Lines like "- Include sections for..."
    ]
    
    for pattern in instruction_lines:
        cleaned_text = re.sub(pattern, "", cleaned_text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove excessive whitespace and clean up formatting
    cleaned_text = re.sub(r'\n\s*\n\s*\n', '\n\n', cleaned_text)  # Max 2 consecutive newlines
    cleaned_text = cleaned_text.strip()
    
    # Log if significant cleaning occurred
    original_length = len(text)
    cleaned_length = len(cleaned_text)
    if original_length - cleaned_length > 100:
        logger.info(f"Cleaned AI thinking traces: removed {original_length - cleaned_length} characters")
    
    return cleaned_text

def add_markdown_paragraph(doc, line):
    """Convert Markdown bold (**text**) into actual bold text
    and add line to Word doc as a paragraph."""
    para = doc.add_paragraph()
    bold_pattern = r"\*\*(.*?)\*\*"
    last_end = 0
    
    for match in re.finditer(bold_pattern, line):
        # Normal text before bold
        if match.start() > last_end:
            para.add_run(line[last_end:match.start()])
        
        # Bold text
        bold_text = match.group(1)
        bold_run = para.add_run(bold_text)
        bold_run.bold = True
        last_end = match.end()
    
    # Remaining text after last bold
    if last_end < len(line):
        para.add_run(line[last_end:])

def generate_docx(response_content, metadata):
    """Generate a DOCX file with proper formatting and page breaks, or RTF if DOCX not available."""
    if DOCX_AVAILABLE:
        doc = Document()
        ts = metadata["timestamp"]
        
        # Title Page
        doc.add_heading('Response Template', level=1)
        doc.add_paragraph(f"Generated: {ts}")
        doc.add_paragraph(f"Source File: {metadata['source_file']}")
        doc.add_paragraph(f"Solicitation: {metadata['solicitation']}")
        doc.add_paragraph(f"Match Score: {metadata['match_score']}")
        doc.add_paragraph(f"Agent: {metadata['agent']}")
        doc.add_paragraph("---")
        
        # Response content, line by line
        for line in response_content.splitlines():
            line = line.strip()
            if not line:
                continue
            
            # Insert page breaks before each section
            if line.lower().startswith("**section 1:"):
                doc.add_page_break()  # Page break before Section 1
                doc.add_heading("Section 1: Human-Readable Summary (For Sender)", level=2)
                continue
            elif line.lower().startswith("**section 2:"):
                doc.add_page_break()
                doc.add_heading("Section 2: Draft Response Template (Formal Email to Government)", level=2)
                continue
            
            # Convert markdown bold (**text**) into real bold runs
            add_markdown_paragraph(doc, line)
        
        return doc
    else:
        # Generate RTF format that can be opened by Word
        return generate_rtf_document(response_content, metadata)

def generate_rtf_document(response_content, metadata):
    """Generate RTF document that can be opened by Microsoft Word."""
    ts = metadata["timestamp"]
    
    # Start RTF document - using \par with space for proper RTF formatting
    rtf_content = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
\f0\fs24 
{\b\fs32 RESPONSE TEMPLATE}\par 
\par 
Generated: """ + ts + r"""\par 
Source File: """ + metadata['source_file'].replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"""\par 
Solicitation: """ + metadata['solicitation'] + r"""\par 
Match Score: """ + metadata['match_score'] + r"""\par 
Agent: """ + metadata['agent'] + r"""\par 
---\par 
\par 
"""
    
    # Process response content
    for line in response_content.splitlines():
        line = line.strip()
        if not line:
            rtf_content += r"\par "
            continue
        
        # Handle section headers
        if line.lower().startswith("**section 1:"):
            rtf_content += r"\page "  # Page break before Section 1
            rtf_content += r"{\b\fs28 SECTION 1: HUMAN-READABLE SUMMARY (FOR SENDER)}\par "
            rtf_content += r"\par "  # Add extra line break
            continue
        elif line.lower().startswith("**section 2:"):
            rtf_content += r"\page "  # Page break
            rtf_content += r"{\b\fs28 SECTION 2: DRAFT RESPONSE TEMPLATE (FORMAL EMAIL TO GOVERNMENT)}\par "
            rtf_content += r"\par "  # Add extra line break
            continue
        
        # Handle subsection headers (like **Opportunity Overview**)
        if line.startswith("**") and line.endswith("**") and len(line) > 4:
            # This is a subsection header
            header_text = line[2:-2]  # Remove ** from both ends
            rtf_content += r"\par {\b " + header_text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"}\par "
            continue
        
        # Convert markdown bold and escape RTF characters
        rtf_line = convert_markdown_to_rtf(line)
        rtf_content += rtf_line + r"\par "
    
    rtf_content += r"}"
    return rtf_content.encode('utf-8')

def convert_markdown_to_rtf(line):
    """Convert markdown bold (**text**) to RTF bold formatting."""
    # Escape RTF special characters first
    line = line.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
    
    # Convert **text** to RTF bold
    bold_pattern = r"\*\*(.*?)\*\*"
    line = re.sub(bold_pattern, r"{\\b \1}", line)
    
    return line

def save_response_to_s3(response_content, json_data, original_key, formats=None):
    """Save the generated response template to S3 in multiple formats.
    Formats determined by OUTPUT_FORMATS environment variable or passed parameter."""
    if formats is None:
        formats = OUTPUT_FORMATS  # Use environment variable
    
    try:
        solicitation_number = json_data.get('solicitationNumber', 'UNKNOWN')
        
        # Extract prefix (parent folder path) from the original key
        prefix = os.path.dirname(original_key)
        if prefix in ["", "."]:
            prefix = ""  # no subfolder
        else:
            prefix = prefix + "/"  # ensure trailing slash
        
        # Common filename stem
        base_filename = os.path.splitext(os.path.basename(original_key))[0]
        output_basename = f"{solicitation_number}_response_template"
        
        # Metadata
        timestamp = datetime.utcnow().isoformat()
        header = f"""Response Template

Generated: {timestamp}
Source: {solicitation_number}
Match Score: {json_data.get('score', 'N/A')}

---

"""
        
        # --- Save as TXT (Default) ---
        if "txt" in formats:
            txt_filename = f"{prefix}{output_basename}.txt"
            full_content = header + response_content
            
            s3_client.put_object(
                Bucket=DESTINATION_BUCKET,
                Key=txt_filename,
                Body=full_content.encode('utf-8'),
                ContentType='text/plain',
                Metadata={
                    'source-file': original_key,
                    'solicitation-number': solicitation_number,
                    'generated-timestamp': timestamp,
                    'match-score': str(json_data.get('score', 'N/A')),
                    'model-used': MODEL_ID
                }
            )
            logger.info(f"Saved TXT response template: {txt_filename}")
        
        # --- Save as DOCX or RTF ---
        if "docx" in formats:
            metadata = {
                "timestamp": timestamp,
                "source_file": original_key,
                "solicitation": solicitation_number,
                "match_score": str(json_data.get('score', 'N/A')),
                "agent": MODEL_ID,
            }
            
            if DOCX_AVAILABLE:
                doc = generate_docx(response_content, metadata)
                
                # Save DOCX in memory
                byte_stream = BytesIO()
                doc.save(byte_stream)
                byte_stream.seek(0)
                
                docx_filename = f"{prefix}{output_basename}.docx"
                s3_client.put_object(
                    Bucket=DESTINATION_BUCKET,
                    Key=docx_filename,
                    Body=byte_stream.getvalue(),
                    ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    Metadata={
                        'source-file': original_key,
                        'solicitation-number': solicitation_number,
                        'generated-timestamp': timestamp,
                        'match-score': str(json_data.get('score', 'N/A')),
                        'model-used': MODEL_ID
                    }
                )
                logger.info(f"Saved DOCX response template with formatting: {docx_filename}")
            else:
                # Generate RTF format instead
                rtf_content = generate_rtf_document(response_content, metadata)
                
                rtf_filename = f"{prefix}{output_basename}.rtf"
                s3_client.put_object(
                    Bucket=DESTINATION_BUCKET,
                    Key=rtf_filename,
                    Body=rtf_content,
                    ContentType="application/rtf",
                    Metadata={
                        'source-file': original_key,
                        'solicitation-number': solicitation_number,
                        'generated-timestamp': timestamp,
                        'match-score': str(json_data.get('score', 'N/A')),
                        'model-used': MODEL_ID
                    }
                )
                logger.info(f"Saved RTF response template (Word-compatible): {rtf_filename}")
            
    except Exception as e:
        logger.error(f"Error saving to S3: {str(e)}")
        raise e