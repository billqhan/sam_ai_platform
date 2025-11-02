"""
Report generation utilities for creating text reports, Word documents, and email templates.
"""

from typing import Dict, Any, Optional
import io
import logging
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    # Handle case where python-docx is not available
    Document = None
    logging.warning("python-docx not available, Word document generation will be disabled")

from template_manager import TemplateManager
import re

logger = logging.getLogger(__name__)

class ReportGenerator:
    """Generates various types of reports from match result data."""
    
    def __init__(self, template_manager: TemplateManager, company_info: Dict[str, str]):
        """
        Initialize report generator.
        
        Args:
            template_manager: Template manager instance
            company_info: Company information dictionary
        """
        self.template_manager = template_manager
        self.company_info = company_info
        self.docx_available = Document is not None
    
    def generate_text_report(self, match_data: Dict[str, Any]) -> str:
        """
        Generate a comprehensive text report from match data.
        
        Args:
            match_data: Match result data from AI processing
            
        Returns:
            str: Formatted text report
        """
        solicitation_id = match_data.get('solicitationNumber') or match_data.get('solicitation_id', 'Unknown')
        logger.info(f"Generating text report: solicitation_id={solicitation_id}")
        
        # Prepare template data
        template_data = self.template_manager.prepare_template_data(match_data, self.company_info)
        
        # Build report sections
        report_sections = []
        
        # Header
        header = self.template_manager.format_template(
            self.template_manager.get_text_template('header'),
            template_data
        )
        report_sections.append(header)
        
        # Opportunity Summary
        opportunity_summary = self.template_manager.format_template(
            self.template_manager.get_text_template('opportunity_summary'),
            template_data
        )
        report_sections.append(opportunity_summary)
        
        # Match Analysis
        match_analysis = self.template_manager.format_template(
            self.template_manager.get_text_template('match_analysis'),
            template_data
        )
        report_sections.append(match_analysis)
        
        # Required Skills (if available)
        if match_data.get('opportunity_required_skills'):
            required_skills = self.template_manager.format_template(
                self.template_manager.get_text_template('required_skills'),
                template_data
            )
            report_sections.append(required_skills)
        
        # Company Skills (if available)
        if match_data.get('company_skills'):
            company_skills = self.template_manager.format_template(
                self.template_manager.get_text_template('company_skills'),
                template_data
            )
            report_sections.append(company_skills)
        
        # Past Performance (if available)
        if match_data.get('past_performance'):
            past_performance = self.template_manager.format_template(
                self.template_manager.get_text_template('past_performance'),
                template_data
            )
            report_sections.append(past_performance)
        
        # Citations (if available)
        if match_data.get('citations'):
            citations = self.template_manager.format_template(
                self.template_manager.get_text_template('citations'),
                template_data
            )
            report_sections.append(citations)
        
        # Footer
        footer = self.template_manager.format_template(
            self.template_manager.get_text_template('footer'),
            template_data
        )
        report_sections.append(footer)
        
        # Combine all sections
        full_report = "\n".join(report_sections)
        
        logger.info(f"Text report generated successfully: length={len(full_report)}")
        
        return full_report
    
    def generate_word_document(self, match_data: Dict[str, Any]) -> bytes:
        """
        Generate a Microsoft Word document from match data.
        
        Args:
            match_data: Match result data from AI processing
            
        Returns:
            bytes: Word document as bytes
        """
        if not self.docx_available:
            logger.warning("python-docx not available, generating placeholder document")
            return self._generate_placeholder_document(match_data)
        
        logger.info(f"Generating Word document: solicitation_id={match_data.get('solicitation_id')}")
        
        # Create new document
        doc = Document()
        
        # Prepare template data
        template_data = self.template_manager.prepare_template_data(match_data, self.company_info)
        
        # Add title
        title = doc.add_heading('SAM Opportunity Match Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add basic information table
        self._add_basic_info_table(doc, template_data)
        
        # Add opportunity summary
        self._add_opportunity_summary(doc, template_data)
        
        # Add match analysis
        self._add_match_analysis(doc, template_data, match_data)
        
        # Add skills comparison if available
        if match_data.get('opportunity_required_skills') or match_data.get('company_skills'):
            self._add_skills_comparison(doc, template_data, match_data)
        
        # Add past performance if available
        if match_data.get('past_performance'):
            self._add_past_performance(doc, template_data)
        
        # Add citations if available
        if match_data.get('citations'):
            self._add_citations(doc, match_data.get('citations'))
        
        # Add footer
        self._add_document_footer(doc, template_data)
        
        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        logger.info("Word document generated successfully")
        
        return doc_bytes.getvalue()
    
    def generate_docx(self, response_content: str, metadata: Dict[str, Any]) -> Document:
        """
        Generate a DOCX file with proper formatting and page breaks from Bedrock Agent response.
        
        Args:
            response_content: Generated response content from Bedrock Agent
            metadata: Metadata dictionary with timestamp, source_file, etc.
            
        Returns:
            Document: Generated DOCX document
        """
        if not self.docx_available:
            logger.warning("python-docx not available, cannot generate DOCX document")
            raise RuntimeError("python-docx library not available")
        
        doc = Document()
        ts = metadata["timestamp"]
        
        # Title Page
        doc.add_heading('Response Template', level=1)
        doc.add_paragraph(f"Generated: {ts}")
        doc.add_paragraph(f"Source File: {metadata['source_file']}")
        doc.add_paragraph(f"Solicitation: {metadata['solicitation']}")
        doc.add_paragraph(f"Match Score: {metadata['match_score']}")
        doc.add_paragraph(f"Agent: {metadata['agent']}")
        doc.add_paragraph("---")
        
        # Response content, line by line
        for line in response_content.splitlines():
            line = line.strip()
            if not line:
                continue
            
            # Insert page breaks before each section
            if line.lower().startswith("**section 1:"):
                doc.add_page_break()  # Page break before Section 1
                doc.add_heading("Section 1: Human-Readable Summary (For Sender)", level=2)
                continue
            elif line.lower().startswith("**section 2:"):
                doc.add_page_break()
                doc.add_heading("Section 2: Draft Response Template (Formal Email to Government)", level=2)
                continue
            
            # Convert markdown bold (**text**) into real bold runs
            self._add_markdown_paragraph(doc, line)
        
        return doc
    
    def _add_markdown_paragraph(self, doc: Document, line: str) -> None:
        """
        Convert Markdown bold (**text**) into actual bold text and add line to Word doc as a paragraph.
        
        Args:
            doc: Document object
            line: Line of text with potential markdown formatting
        """
        para = doc.add_paragraph()
        bold_pattern = r"\*\*(.*?)\*\*"
        last_end = 0
        
        for match in re.finditer(bold_pattern, line):
            # Normal text before bold
            if match.start() > last_end:
                para.add_run(line[last_end:match.start()])
            
            # Bold text
            bold_text = match.group(1)
            bold_run = para.add_run(bold_text)
            bold_run.bold = True
            last_end = match.end()
        
        # Remaining text after last bold
        if last_end < len(line):
            para.add_run(line[last_end:])
    
    def generate_email_template(self, match_data: Dict[str, Any]) -> str:
        """
        Generate an email template for POC outreach.
        
        Args:
            match_data: Match result data from AI processing
            
        Returns:
            str: Formatted email template
        """
        logger.info(f"Generating email template: solicitation_id={match_data.get('solicitation_id')}")
        
        # Prepare template data
        template_data = self.template_manager.prepare_template_data(match_data, self.company_info)
        
        # Generate subject line
        subject = self.template_manager.format_template(
            self.template_manager.get_email_template('subject'),
            template_data
        )
        
        # Choose email body based on match status
        is_match = match_data.get('is_match', False)
        body_template = 'body' if is_match else 'no_match_body'
        
        body = self.template_manager.format_template(
            self.template_manager.get_email_template(body_template),
            template_data
        )
        
        # Combine subject and body
        email_template = f"Subject: {subject}\n\n{body}"
        
        logger.info(f"Email template generated successfully: is_match={is_match}")
        
        return email_template
    
    def _generate_placeholder_document(self, match_data: Dict[str, Any]) -> bytes:
        """Generate a Rich Text Format document that can be opened by Word."""
        # Prepare template data
        template_data = self.template_manager.prepare_template_data(match_data, self.company_info)
        
        # Generate RTF document that Word can open
        rtf_content = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
\f0\fs24 
{\b\fs32 SAM OPPORTUNITY MATCH REPORT\par}
\par
{\b === BASIC INFORMATION ===\par}
{\b Solicitation ID:} """ + template_data.get('solicitation_id', 'Unknown') + r"""\par
{\b Title:} """ + template_data.get('title', 'Unknown').replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"""\par
{\b Match Score:} """ + f"{template_data.get('match_score', 0.0):.2f} ({template_data.get('match_percentage', 0.0):.1f}%)" + r"""\par
{\b Match Status:} """ + template_data.get('match_status', 'Unknown') + r"""\par
{\b Generated:} """ + template_data.get('timestamp', 'Unknown') + r"""\par
\par
{\b === OPPORTUNITY SUMMARY ===\par}
{\b Title:} """ + template_data.get('title', 'Unknown').replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"""\par
{\b Deadline:} """ + template_data.get('deadline', 'Not specified') + r"""\par
{\b Value:} """ + template_data.get('value', 'Not specified') + r"""\par
\par
{\b === MATCH ANALYSIS ===\par}
{\b Match Score:} """ + f"{template_data.get('match_score', 0.0):.2f}" + r"""\par
{\b Match Status:} """ + template_data.get('match_status', 'Unknown') + r"""\par
\par
{\b Rationale:\par}
""" + match_data.get('rationale', 'No rationale provided').replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}') + r"""\par
\par
{\b === OPPORTUNITY REQUIRED SKILLS ===\par}
""" + self._format_skills_for_rtf(match_data.get('opportunity_required_skills', [])) + r"""\par
\par
{\b === COMPANY MATCHING SKILLS ===\par}
""" + self._format_skills_for_rtf(match_data.get('company_skills', [])) + r"""\par
\par
{\b === PAST PERFORMANCE ===\par}
""" + self._format_skills_for_rtf(match_data.get('past_performance', [])) + r"""\par
\par
{\b === SUPPORTING CITATIONS ===\par}
""" + self._format_citations_for_rtf(match_data.get('citations', [])) + r"""\par
\par
{\b === REPORT INFORMATION ===\par}
Report generated by AI RFP Response Agent\par
{\b Company:} """ + template_data.get('company_name', 'Your Company') + r"""\par
{\b Contact:} """ + template_data.get('company_contact', 'contact@yourcompany.com') + r"""\par
{\b Generated:} """ + template_data.get('timestamp', 'Unknown') + r"""\par
\par
{\i Note: This RTF document can be opened and edited in Microsoft Word.\par}
}"""
        
        return rtf_content.encode('utf-8')
    
    def _format_skills_for_document(self, skills: list) -> str:
        """Format skills list for document display."""
        if not skills:
            return "None specified"
        return "\n".join(f"• {skill}" for skill in skills)
    
    def _format_citations_for_document(self, citations: list) -> str:
        """Format citations for document display."""
        if not citations:
            return "No citations available"
        
        formatted_citations = []
        for i, citation in enumerate(citations, 1):
            doc_title = citation.get('document_title', 'Unknown Document')
            section = citation.get('section_or_page', 'Unknown Section')
            excerpt = citation.get('excerpt', 'No excerpt available')
            
            formatted_citation = f"{i}. {doc_title} - {section}\n   \"{excerpt}\""
            formatted_citations.append(formatted_citation)
        
        return "\n\n".join(formatted_citations)
    
    def _format_skills_for_rtf(self, skills: list) -> str:
        """Format skills list for RTF document."""
        if not skills:
            return "None specified"
        
        rtf_skills = []
        for skill in skills:
            # Escape RTF special characters
            escaped_skill = skill.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            rtf_skills.append(f"\\bullet {escaped_skill}\\par")
        
        return "\n".join(rtf_skills)
    
    def _format_citations_for_rtf(self, citations: list) -> str:
        """Format citations for RTF document."""
        if not citations:
            return "No citations available"
        
        rtf_citations = []
        for i, citation in enumerate(citations, 1):
            doc_title = citation.get('document_title', 'Unknown Document')
            section = citation.get('section_or_page', 'Unknown Section')
            excerpt = citation.get('excerpt', 'No excerpt available')
            
            # Escape RTF special characters
            escaped_title = doc_title.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            escaped_section = section.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            escaped_excerpt = excerpt.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
            
            rtf_citation = f"{i}. {escaped_title} - {escaped_section}\\par   \\\"{escaped_excerpt}\\\"\\par"
            rtf_citations.append(rtf_citation)
        
        return "\\par\n".join(rtf_citations)
    
    def _add_basic_info_table(self, doc, template_data: Dict[str, Any]) -> None:
        """Add basic information table to Word document."""
        doc.add_heading('Basic Information', level=1)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        
        # Add data to table
        cells = table.rows[0].cells
        cells[0].text = 'Solicitation ID'
        cells[1].text = template_data.get('solicitation_id', 'Unknown')
        
        cells = table.rows[1].cells
        cells[0].text = 'Match Score'
        cells[1].text = f"{template_data.get('match_score', 0.0):.2f} ({template_data.get('match_percentage', 0.0):.1f}%)"
        
        cells = table.rows[2].cells
        cells[0].text = 'Match Status'
        cells[1].text = template_data.get('match_status', 'Unknown')
        
        cells = table.rows[3].cells
        cells[0].text = 'Generated'
        cells[1].text = template_data.get('timestamp', 'Unknown')
    
    def _add_opportunity_summary(self, doc, template_data: Dict[str, Any]) -> None:
        """Add opportunity summary section to Word document."""
        doc.add_heading('Opportunity Summary', level=1)
        
        doc.add_paragraph(f"Title: {template_data.get('title', 'Unknown')}")
        doc.add_paragraph(f"Value: {template_data.get('value', 'Not specified')}")
        doc.add_paragraph(f"Deadline: {template_data.get('deadline', 'Not specified')}")
        doc.add_paragraph(f"NAICS Codes: {template_data.get('naics_codes', 'None specified')}")
    
    def _add_match_analysis(self, doc, template_data: Dict[str, Any], match_data: Dict[str, Any]) -> None:
        """Add match analysis section to Word document."""
        doc.add_heading('Match Analysis', level=1)
        
        doc.add_paragraph(f"Match Score: {template_data.get('match_score', 0.0):.2f} ({template_data.get('match_percentage', 0.0):.1f}%)")
        doc.add_paragraph(f"Match Status: {template_data.get('match_status', 'Unknown')}")
        
        doc.add_heading('Rationale', level=2)
        rationale = match_data.get('rationale', 'No rationale provided')
        doc.add_paragraph(rationale)
    
    def _add_skills_comparison(self, doc, template_data: Dict[str, Any], match_data: Dict[str, Any]) -> None:
        """Add skills comparison section to Word document."""
        doc.add_heading('Skills Comparison', level=1)
        
        # Required skills
        required_skills = match_data.get('opportunity_required_skills', [])
        if required_skills:
            doc.add_heading('Opportunity Required Skills', level=2)
            for skill in required_skills:
                p = doc.add_paragraph(skill)
                p.style = 'List Bullet'
        
        # Company skills
        company_skills = match_data.get('company_skills', [])
        if company_skills:
            doc.add_heading('Company Matching Skills', level=2)
            for skill in company_skills:
                p = doc.add_paragraph(skill)
                p.style = 'List Bullet'
    
    def _add_past_performance(self, doc, template_data: Dict[str, Any]) -> None:
        """Add past performance section to Word document."""
        doc.add_heading('Relevant Past Performance', level=1)
        
        past_performance = template_data.get('past_performance', 'None specified')
        doc.add_paragraph(past_performance)
    
    def _add_citations(self, doc, citations: list) -> None:
        """Add citations section to Word document."""
        doc.add_heading('Supporting Citations', level=1)
        
        for i, citation in enumerate(citations, 1):
            doc_title = citation.get('document_title', 'Unknown Document')
            section = citation.get('section_or_page', 'Unknown Section')
            excerpt = citation.get('excerpt', 'No excerpt available')
            
            doc.add_heading(f'Citation {i}', level=2)
            doc.add_paragraph(f"Document: {doc_title}")
            doc.add_paragraph(f"Section: {section}")
            doc.add_paragraph(f"Excerpt: \"{excerpt}\"")
    
    def _add_document_footer(self, doc, template_data: Dict[str, Any]) -> None:
        """Add footer information to Word document."""
        doc.add_page_break()
        
        footer_heading = doc.add_heading('Report Information', level=1)
        footer_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report generated by AI RFP Response Agent")
        doc.add_paragraph(f"Company: {template_data.get('company_name', 'Your Company')}")
        doc.add_paragraph(f"Contact: {template_data.get('company_contact', 'contact@yourcompany.com')}")
        doc.add_paragraph(f"Generated: {template_data.get('timestamp', 'Unknown')}")